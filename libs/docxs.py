#!/usr/bin/python3.10


################################################################################
#                                                                              #
#                                                                              #
#                   GNU AFFERO GENERAL PUBLIC LICENSE                          #
#                       Version 3, 19 November 2007                            #
#                                                                              #
#    This programms aims at sanitizing .docx and .docm documents.              #
#    Copyright (C) 2022 Gavroche                                               #
#                                                                              #
#    This program is free software: you can redistribute it and/or modify      #
#    it under the terms of the GNU Affero General Public License as published  #
#    by the Free Software Foundation, either version 3 of the License, or      #
#    (at your option) any later version.                                       #
#                                                                              #
#     This program is distributed in the hope that it will be useful,          #
#    but WITHOUT ANY WARRANTY; without even the implied warranty of            #
#    MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the             #
#    GNU Affero General Public License for more details.                       #
#                                                                              #
#    You should have received a copy of the GNU Affero General Public License  #
#    along with this program.  If not, see <https://www.gnu.org/licenses/>.    #
#                                                                              #
#                                                                              #
################################################################################


import os

import xml.etree.ElementTree as ET

from docx import Document
from docx.shared import Inches
from zipfile import ZipFile
from PIL.Image import *
from io import BytesIO


################################################################################


# Define XML namespace (Microsoft)
xmlns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        }


################################################################################


# Get document layout
def get_layout(path):

    layout = [path.split("/")[-1]]

    # Use doc as zip file
    unzip = ZipFile(path)

    # Parse xml
    root = ET.fromstring(unzip.read('word/document.xml'))
    body = root.find('w:body', xmlns)

    for section in body:
        if section.tag[-1:] == 'p':
            if section.findall('.//w:t', xmlns) != [] or section.find('.//pic:cNvPr', xmlns) is not None:
                layout.append('p')
            else:
                pass

        elif section.tag[-3:] == 'tbl':
            layout.append('tbl')

    return layout


################################################################################


# Get content from a section
def get_content(section, id):

    ctnt = []

    # Get all text nodes from the section. Empty list if none
    texts = section.findall('.//w:t', xmlns)

    # Get image node from the section. None type if none
    image = section.find('.//pic:cNvPr', xmlns)

    if texts != []:
        # Joins text parts if it exists
        ctnt.append({'text': ''.join([t.text for t in texts])})

    if image is not None:
        # Gets image id if it exists
        ctnt.append({'img': id, 'w': section.find('.//wp:extent', xmlns).get('cx') , 'h': section.find('.//wp:extent', xmlns).get('cy')})
        id += 1

    return ctnt, id


################################################################################


def get_sections(path):

    content = ([], [])
    id, jd = 0, 0

    # Use doc as zip file
    unzip = ZipFile(path)

    # Parse xml
    root = ET.fromstring(unzip.read('word/document.xml'))
    body = root.find('w:body', xmlns)
    sections = body.findall('w:p', xmlns)

    # Retrieve p sections content
    for section in sections:

        ctnt, id = get_content(section, id)

        if ctnt != []: content[0].append(ctnt)


    # Retrieve tables
    tables = body.findall('w:tbl', xmlns)

    for table in tables:

        # Gets table data if it exists
        cases = {'tbl': jd, 'cases': []}
        rows = table.findall('.//w:tr', xmlns)
        for row in rows:
            columns = row.findall('.//w:tc', xmlns)
            for column in columns:
                for section in column.findall('w:p', xmlns):
                    ctnt, id = get_content(section, id)
                    if ctnt != []:
                        case = (ctnt, rows.index(row), columns.index(column))
                        cases['cases'].append(case)

        content[1].append(cases)
        jd += 1

    return content


################################################################################


def get_imgs(path):

    img_list = []

    unzip = ZipFile(path)

    for file in unzip.namelist():
        if file.startswith('word/media/'):

            # Extracting source image
            name = file.split("/")[-1]
            img = open(BytesIO(unzip.read(file)))
            img.save('Outputs/' + name)

            # Sanitizing image
            ext_img('Outputs/' + name, 'Outputs/out_' + name)

            # Adding sanitized image to image list
            img_list.append('Outputs/out_' + name)


    return img_list


################################################################################


def recompose(layout, content, img_list):

    # New document creation
    doc = Document()

    # Document reconstitution
    for item in layout:
        if item == 'p':
            ctnt = content[0].pop(0)

            for sec in ctnt:

                if 'text' in sec:
                    doc.add_paragraph(sec['text'])

                elif 'img' in sec:
                    path = img_list[int(sec['img'])]

                    # Unit conversion
                    w = int(sec['w'])/914400
                    h = int(sec['h'])/914400

                    doc.add_picture(path, width=Inches(w), height=Inches(h))

                    os.remove(path)


        elif item == 'tbl':
            tbl = content[1].pop(0)
            rws = max([c[1] for c in tbl['cases']]) + 1
            cls = max([c[2] for c in tbl['cases']]) + 1

            table = doc.add_table(rows=rws, cols=cls, style="Table Grid")

            for case in tbl['cases']:
                ct = case[0]
                ro = case[1]
                co = case[2]

                for sec in ct:

                    # If case contains text
                    if 'text' in sec:
                        table.rows[ro].cells[co].paragraphs[0].text += sec['text']

                    #If case contains an image
                    if 'img' in sec:

                        path = img_list[int(sec['img'])]

                        # Unit conversion
                        w = int(sec['w'])/914400
                        h = int(sec['h'])/914400

                        table.rows[ro].cells[co].add_paragraph()
                        table.rows[ro].cells[co].paragraphs[-1].add_run().add_picture(path, width=Inches(w), height=Inches(h))

                        os.remove(path)


        doc.save('Outputs/out_' + layout[0].split("/")[-1].split(".")[0] + '.docx')


################################################################################


def sanitiz(path):

    info('Sanitizing ' + path.split("/")[-1])

    # Fetching data from source
    try:
        l = get_layout(path)
        c = get_sections(path)
        i = get_imgs(path)

    except:
        fail('Data extraction failed')
        return False, ''


    info('Creating new document')

    # Creating sanitized document
    try:
        recompose(l, c, i)

    except:
        fail('Document recomposition failed')
        return False, ''

    success('Document sanitized successfully.')
    return True, 'Outputs/out_' + path.split("/")[-1].split(".")[0] + '.docx'


################################################################################


if __name__ == "__main__":
    print('Please run main.py or read software documentation')

    exit()

else:
    from .imgs import *
    from .prints import *
