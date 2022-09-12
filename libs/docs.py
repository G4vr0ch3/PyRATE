#!/usr/bin/python3.10


################################################################################
#                                                                              #
#                                                                              #
#                   GNU AFFERO GENERAL PUBLIC LICENSE                          #
#                       Version 3, 19 November 2007                            #
#                                                                              #
#    This programms aims at sanitizing .doc documents.                         #
#    Copyright (C) 2022 Gavroche                                               #
#                                                                              #
#    This program is free software: you can redistribute it and/or modify      #
#    it under the terms of the GNU Affero General Public License as published  #
#    by the Free Software Foundation, either version 3 of the License, or      #
#    (at your option) any later version.                                       #
#                                                                              #
#     This program is distributed in the hope that it will be useful,          #
#    but WITHOUT ANY WARRANTY; without even the implied warranty of            #
#    MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the             #
#    GNU Affero General Public License for more details.                       #
#                                                                              #
#    You should have received a copy of the GNU Affero General Public License  #
#    along with this program.  If not, see <https://www.gnu.org/licenses/>.    #
#                                                                              #
#                                                                              #
################################################################################


# This software relies partly on the ANTIWORD software developed by Adri van Os.
# For more information regarding this sofwtare : visit http://www.winfield.demon.nl/


################################################################################


import subprocess
import os

import PIL.Image as PI

from docx import Document
from docx.shared import Inches


################################################################################


def get_layout(path):

    # Execute os command
    pipe = subprocess.Popen(
        'antiword "{}"'.format(path), shell=True,
        stdout=subprocess.PIPE, stderr=subprocess.PIPE,
    )

    # Wait for subprocess completion
    pipe.wait()

    # If pipe fails
    if pipe.returncode != 0:
        return False

    # Read pipe result
    raw = pipe.stdout.read()

    raw_layout = raw.split(b'\n\n')
    layout = []

    # Build layout
    for item in raw_layout:

        if item == b'':
            pass

        elif item == b'[pic]':
            layout.append('img')

        elif item.startswith(b'|'):
            layout.append('tbl')

        else:
            layout.append('text')

    return layout


################################################################################


def build_text(doc, raw_txt, index):

    # Text from bytes
    beautify = raw_txt[index].decode().strip()

    doc.add_paragraph(beautify)


################################################################################


def build_img(doc, img_list):

    # Retrieve image path from image list
    try:
        img_path = img_list.pop(0)['path']

        doc.add_picture(img_path, width=Inches(5))

        # Remove image once used
        os.remove(img_path)

    # If the image is missing
    except:
        build_text(doc, ['Missing Image'], 0)
        warning('Missing Image')


################################################################################


def build_tbl(doc, raw_txt, index, img_list):

    cases = []

    raw_cases = raw_txt[index].split(b'|')

    # Retrieve column count
    cols = raw_cases.index(b'\n') - 1

    # Get each case's content
    for case in raw_cases:
        if case != b'' and case != b'\n':
            cases.append(case.decode().strip())

    # Calculate row count
    rows = len(cases)//cols

    # Create table
    table = doc.add_table(rows=rows, cols=cols, style="Table Grid")

    for r in range(rows):
        for c in range(cols):

            case = cases[r*cols + c]

            # Handle image case
            if case == '[pic]':

                # If there is a discrepency between the image count from the layout and the actually retrieved images
                try:
                    img_path = img_list.pop(0)['path']

                    table.rows[r].cells[c].paragraphs[0].add_run().add_picture(img_path, width = Inches(3))

                    os.remove(img_path)

                except:
                    pass

            # Text otherwise
            else:
                table.rows[r].cells[c].paragraphs[0].text += case


################################################################################


# Create sanitized document
def recompose(path, layout, img_list):

    raw_txt = []

    # Execute os command
    pipe = subprocess.Popen(
        'antiword {}'.format(path), shell=True,
        stdout=subprocess.PIPE, stderr=subprocess.PIPE,
    )

    # Wait for subprocess completion
    pipe.wait()

    # If pipe fails
    if pipe.returncode != 0:
        return False

    # Read pipe result
    raw = pipe.stdout.read().split(b'\n\n')

    for item in raw:
        if item != b'': raw_txt.append(item)

    doc = Document()

    for index in range(len(layout)):

        # Handle regular text paragraph
        if layout[index] == 'text':
            build_text(doc, raw_txt, index)

        # Handle images
        elif layout[index] == 'img':
            build_img(doc, img_list)

        # Handle tables
        elif layout[index] == 'tbl':
            build_tbl(doc, raw_txt, index, img_list)

    doc.save('Outputs/out_' + path.split("/")[-1].split(".")[0] + '.docx')


################################################################################


def sanitiz(path):

    info('Sanitizing ' + path.split("/")[-1])

    try:
        # Fetching data from source
        layout = get_layout(path)
        img_list = common_bin.img_parser(path)


        try:
            if len(img_list) != layout.count('img'): warning('Image count from layout and retrieved image count differ. The document images might not appear as attended. This is quite common and not fatal.')
        except:
            pass


    except:
        fail('Data extraction failed')
        return False, ''


    info('Creating new document')

    try:
        # Creating sanitized document
        recompose(path, layout, img_list)

    except:
        fail('Document recomposition failed')
        return False, ''

    success('Document sanitized successfully.')
    return True, 'Outputs/out_' + path.split("/")[-1].split(".")[0] + '.docx'


################################################################################


if __name__ == "__main__":
    print('Please run main.py or read software documentation')

    exit()

else:
    from .prints import *
    from .imgs import *
    from . import common_bin
